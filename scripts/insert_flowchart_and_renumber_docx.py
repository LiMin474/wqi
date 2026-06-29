from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


BASE_DOCX = Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\paper_sections_3_5_updated.docx")
FLOWCHART = Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\results\method_framework_polished.png")
OUT_DOCX = Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\AI论文初稿_第3_5章修改_含流程图.docx")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = paragraph.insert_paragraph_before(text, style=style)
    paragraph._element.addprevious(new_p._element)
    return new_p


def find_para(doc: Document, contains: str):
    for para in doc.paragraphs:
        if contains in para.text:
            return para
    raise ValueError(f"Paragraph containing {contains!r} not found")


def renumber_text(text: str) -> str:
    replacements = [
        ("图5", "图6"),
        ("图 5", "图6"),
        ("图4", "图5"),
        ("图 4", "图5"),
        ("图3", "图4"),
        ("图 3", "图4"),
        ("图2", "图3"),
        ("图 2", "图3"),
        ("图1", "图2"),
        ("图 1", "图2"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def main():
    doc = Document(str(BASE_DOCX))

    # Insert polished flowchart in Section 2 after the end of 2.3
    anchor = find_para(doc, "该方法具有以下优势")
    caption_para = anchor.insert_paragraph_before("图1. 所提多进化算法优化与加权集成框架示意图", style="Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pic_para = anchor.insert_paragraph_before("", style="Normal")
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(FLOWCHART), width=Inches(6.7))

    # Add a short lead-in sentence before figure
    lead = caption_para.insert_paragraph_before(
        "为便于整体说明本文方法流程，图1展示了从数据准备、六种进化算法优化 ANN，到 WeightedAvg 集成与最终评估的完整框架。",
        style="Normal",
    )

    # Renumber later figure references/captions in sections 3-5
    start_para = find_para(doc, "3 实验")
    started = False
    for para in doc.paragraphs:
        if para == start_para:
            started = True
        if not started:
            continue
        if para.text:
            para.text = renumber_text(para.text)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
