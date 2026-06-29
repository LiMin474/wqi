from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


NEW_IMAGE = Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\results\method_framework_polished.png")
DOCX_FILES = [
    Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\final_paper_checked.docx"),
    Path(r"D:\study\project\water-quality\Stacking-WQI-Hyperopt\AI论文初稿_第3_5章修改_含流程图_终校版.docx"),
]


def replace_first_image_after_caption(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    target_index = None
    for i, para in enumerate(doc.paragraphs):
        if "图1. 所提多进化算法优化与加权集成框架示意图" in (para.text or ""):
            target_index = i + 1
            break

    if target_index is None or target_index >= len(doc.paragraphs):
        raise ValueError(f"Flowchart caption not found in {docx_path}")

    para = doc.paragraphs[target_index]
    runs = para.runs
    if not runs:
        raise ValueError(f"No image run found after caption in {docx_path}")

    replaced = False
    new_blob = NEW_IMAGE.read_bytes()
    for run in runs:
        blips = run._element.xpath('.//a:blip')
        if not blips:
            continue
        rid = blips[0].get(qn('r:embed'))
        if not rid:
            continue
        image_part = doc.part.related_parts[rid]
        image_part._blob = new_blob
        replaced = True
        break

    if not replaced:
        raise ValueError(f"No embedded image found after caption in {docx_path}")

    doc.save(str(docx_path))


def main() -> None:
    for path in DOCX_FILES:
        replace_first_image_after_caption(path)
        print(f"Updated flowchart in: {path}")


if __name__ == "__main__":
    main()
