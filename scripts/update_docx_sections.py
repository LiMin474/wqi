from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


TABLE_NUMBER_MAP = {
    "表 1": "表 3",
    "表1": "表3",
    "表 2": "表 4",
    "表2": "表4",
    "表 3": "表 5",
    "表3": "表5",
    "表 4": "表 6",
    "表4": "表6",
}


def normalize_text(text: str) -> str:
    replacements = {
        "$R^2_{CV}$": "R²CV",
        "$R^2$": "R²",
        "Wilcoxon $p$": "Wilcoxon p",
        "Friedman $p$": "Friedman p",
        "$p$": "p",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)

    text = re.sub(r"\$p\s*=\s*([0-9.]+)\$", r"p = \1", text)
    text = text.replace("$", "")

    for src, dst in TABLE_NUMBER_MAP.items():
        text = text.replace(src, dst)

    text = text.replace("图 1", "图1")
    text = text.replace("图 2", "图2")
    text = text.replace("图 3", "图3")
    text = text.replace("图 4", "图4")
    text = text.replace("图 5", "图5")

    text = text.replace("（$p = ", "(p = ").replace("$）", ")")
    return text.strip()


def parse_markdown(md_path: Path) -> list[dict]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append({"type": "heading", "level": 3, "text": normalize_text(stripped[4:])})
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append({"type": "heading", "level": 2, "text": normalize_text(stripped[3:])})
            i += 1
            continue
        if stripped.startswith("# "):
            blocks.append({"type": "heading", "level": 1, "text": normalize_text(stripped[2:])})
            i += 1
            continue

        if stripped.startswith("**表"):
            caption = stripped.strip("*")
            caption = normalize_text(caption)
            caption = re.sub(r"^表\s*([0-9]+)\s*", r"表\1. ", caption)
            blocks.append({"type": "table_caption", "text": caption})
            i += 1
            continue

        if stripped.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if not m:
                raise ValueError(f"Invalid image line: {stripped}")
            caption = normalize_text(m.group(1))
            caption = re.sub(r"^图\s*([0-9]+)\s*", r"图\1. ", caption)
            blocks.append({"type": "image", "caption": caption, "path": m.group(2)})
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for idx, tbl_line in enumerate(table_lines):
                if idx == 1 and re.match(r"^\|[:\-\s|]+\|?$", tbl_line):
                    continue
                cells = [normalize_text(c.strip()) for c in tbl_line.strip("|").split("|")]
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "![", "|", "**表")):
                break
            para_lines.append(nxt)
            i += 1
        text = normalize_text(" ".join(para_lines))
        blocks.append({"type": "paragraph", "text": text})

    return blocks


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact_text:
            return para
    raise ValueError(f"Paragraph not found: {exact_text}")


def body_child_index(doc: Document, para: Paragraph) -> int:
    children = list(doc.element.body.iterchildren())
    for idx, child in enumerate(children):
        if child == para._element:
            return idx
    raise ValueError("Body child index not found")


def remove_body_range(doc: Document, start_idx: int, end_idx: int) -> None:
    children = list(doc.element.body.iterchildren())
    for child in children[start_idx:end_idx]:
        doc.element.body.remove(child)


def insert_paragraph_before(anchor: Paragraph, text: str, style: str = "Normal", center: bool = False) -> Paragraph:
    para = anchor.insert_paragraph_before(text, style=style)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def insert_table_before(anchor: Paragraph, doc: Document, rows: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    anchor._element.addprevious(table._element)
    return table


def insert_image_before(anchor: Paragraph, doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")
    doc.add_picture(str(image_path), width=Inches(6.2))
    pic_para = doc.paragraphs[-1]
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor._element.addprevious(pic_para._element)
    insert_paragraph_before(anchor, caption, style="Normal", center=True)


def update_docx(input_docx: Path, md_path: Path, output_docx: Path) -> None:
    doc = Document(str(input_docx))
    blocks = parse_markdown(md_path)

    start_para = find_paragraph(doc, "3 实验")
    appendix_para = find_paragraph(doc, "附录A：各数据集特征列表")

    start_idx = body_child_index(doc, start_para)
    appendix_idx = body_child_index(doc, appendix_para)
    remove_body_range(doc, start_idx, appendix_idx)

    anchor = find_paragraph(doc, "附录A：各数据集特征列表")

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            style = "Heading 2" if block["level"] == 1 else "Heading 3"
            insert_paragraph_before(anchor, block["text"], style=style)
        elif btype == "paragraph":
            insert_paragraph_before(anchor, block["text"], style="Normal")
        elif btype == "table_caption":
            insert_paragraph_before(anchor, block["text"], style="Normal")
        elif btype == "table":
            insert_table_before(anchor, doc, block["rows"])
        elif btype == "image":
            insert_image_before(anchor, doc, Path(block["path"]), block["caption"])
        else:
            raise ValueError(f"Unsupported block type: {btype}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def main() -> int:
    if len(sys.argv) != 4:
        print("Usage: update_docx_sections.py <input_docx> <markdown> <output_docx>")
        return 1

    input_docx = Path(sys.argv[1])
    md_path = Path(sys.argv[2])
    output_docx = Path(sys.argv[3])
    update_docx(input_docx, md_path, output_docx)
    print(f"Saved updated document to: {output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
